from __future__ import annotations

import asyncio
from unittest.mock import AsyncMock, MagicMock

import pytest
from docx import Document

from draftlens_api.domain.enums import IssueCategory, IssueSeverity
from draftlens_api.domain.models import ConflictSet, DocumentBlock, EvidenceSource
from draftlens_api.engine.arbitration_engine import ArbitrationEngine
from draftlens_api.engine.pipeline_evidence import rank_evidence_for_block
from draftlens_api.evidence.retriever import EvidenceRetriever
from draftlens_api.evidence.types import EvidenceChunk
from draftlens_api.routing.agent_assignment import ResolvedModel
from draftlens_api.routing.model_registry import ModelRegistry
from draftlens_api.utils.run_splitting import split_indices_for_runs


def test_rank_evidence_for_block_orders_by_score() -> None:
    block = DocumentBlock(block_id="b1", char_start=0, char_end=20, text="alpha beta gamma particle physics")
    src_hi = EvidenceSource(label="hi", excerpt="alpha beta gamma particle physics details", kind="supporting")
    src_lo = EvidenceSource(label="lo", excerpt="unrelated cooking recipe zucchini", kind="supporting")
    out = rank_evidence_for_block(block, [src_lo, src_hi], max_excerpts=2, max_chars_each=900)
    assert len(out) == 2
    assert "[hi]" in out[0]


def test_evidence_retriever_rank_for_block_jaccard_fallback() -> None:
    block = DocumentBlock(block_id="b1", char_start=0, char_end=10, text="quantum field theory introduction")
    chunks = [
        EvidenceChunk(
            chunk_id="c1",
            source_label="s1",
            text="quantum field theory introduction and operators",
        ),
        EvidenceChunk(
            chunk_id="c2",
            source_label="s2",
            text="baking bread fermentation temperatures",
        ),
    ]
    r = EvidenceRetriever(chunks=chunks, index=None, has_supporting_files=True, partial_file=False)
    rr = r.rank_for_block(block, top_k=2, max_chars_each=500)
    assert rr.excerpts
    assert rr.excerpts[0].chunk_id == "c1"


def test_run_splitting_batches() -> None:
    assert split_indices_for_runs(10, 3) == [(0, 3), (3, 6), (6, 9), (9, 10)]
    with pytest.raises(ValueError):
        split_indices_for_runs(5, 0)


def test_issue_comment_body_includes_sections() -> None:
    from draftlens_api.artifacts import render as render_mod

    body = render_mod._issue_comment_body(
        {
            "category": "grammar",
            "severity": "minor",
            "title": "Comma splice",
            "explanation": "Two independent clauses joined by a comma.",
            "evidence_basis": "",
            "suggested_fix": "Use a semicolon or coordinating conjunction.",
            "confidence": 0.8,
            "source_agents": ["gpt_reviewer"],
            "accuracy_posture": "",
        }
    )
    assert "Title: Comma splice" in body
    assert "Suggested fix:" in body


def test_docx_split_run_splits_text() -> None:
    from draftlens_api.artifacts.docx_run_utils import _split_run

    doc = Document()
    p = doc.add_paragraph()
    run = p.add_run("abcdef")
    a, b = _split_run(run, 3)
    assert a.text == "abc"
    assert b.text == "def"


def test_arbitration_engine_applies_normalization(monkeypatch: pytest.MonkeyPatch) -> None:
    registry = MagicMock(spec=ModelRegistry)
    adapter = MagicMock()
    adapter.complete_json_with_retry = AsyncMock(
        return_value=(
            {
                "executive_summary": "Done",
                "issues": [
                    {
                        "block_id": "bx",
                        "span_text": "x",
                        "char_start": 0,
                        "char_end": 1,
                        "category": "grammar",
                        "severity": "nit",
                        "title": "typo",
                    }
                ],
                "proposed_edits": [],
                "resolved_conflicts": [],
            },
            None,
        )
    )
    registry.adapter.return_value = adapter
    registry.assignment.return_value = MagicMock(arbiter=ResolvedModel(provider="openai", model_id="gpt-test"))

    eng = ArbitrationEngine(registry)  # type: ignore[arg-type]
    decision = asyncio.run(
        eng.run(
            arbiter=ResolvedModel(provider="openai", model_id="gpt-test"),
            debate_digest="d",
            conflicts=[ConflictSet(topic="t", positions={"a": "1", "b": "2"})],
            document_excerpt="hello",
            output_mode="review",
            do_not_change="",
            context="",
        )
    )
    assert decision.issues[0].category == IssueCategory.grammar
    assert decision.issues[0].severity == IssueSeverity.nit
