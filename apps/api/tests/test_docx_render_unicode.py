"""DOCX render helpers preserve Unicode / complex-script fonts where possible."""

from __future__ import annotations

from docx import Document

from draftlens_api.artifacts import docx_run_utils as dx
from draftlens_api.artifacts import render as artifact_render


def test_set_paragraph_text_unicode_safe_preserves_single_run() -> None:
    doc = Document()
    p = doc.add_paragraph()
    p.add_run("模板")
    dx.set_paragraph_text_unicode_safe(p, "中文混合 English text")
    assert p.text == "中文混合 English text"


def test_apply_edits_replaces_without_assigning_paragraph_text() -> None:
    doc = Document()
    doc.add_paragraph("Hello 世界")
    edits = [{"before": "世界", "after": "朋友"}]
    n = artifact_render._apply_edits_to_docx(doc, edits)
    assert n >= 1
    flat, _ = dx.flatten_doc_paragraphs(doc)
    assert "朋友" in flat
    assert "世界" not in flat
