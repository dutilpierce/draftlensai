from __future__ import annotations

import json
import logging
import re
import shutil
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

from docx import Document
from docx.shared import Pt

from draftlens_api.artifacts import docx_run_utils as dx
from draftlens_api.artifacts.disclaimers import build_disclaimer_bundle
from draftlens_api.domain.models import DisclaimerBundle

logger = logging.getLogger(__name__)


def _now_stamp() -> str:
    return datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")


def _locked_phrases(do_not_change: str | None) -> list[str]:
    if not do_not_change or not do_not_change.strip():
        return []
    parts = re.split(r"[\n;,]+", do_not_change)
    return sorted({p.strip() for p in parts if len(p.strip()) > 3}, key=len, reverse=True)


def _partition_edits_for_do_not_change(
    edits: list[dict[str, Any]],
    phrases: list[str],
) -> tuple[list[dict[str, Any]], list[dict[str, Any]]]:
    """Edits that would remove or alter a locked phrase from `before` to `after` are held back."""
    if not phrases:
        return edits, []
    safe: list[dict[str, Any]] = []
    held: list[dict[str, Any]] = []
    for e in edits:
        before = str(e.get("before", ""))
        after = str(e.get("after", ""))
        bl, al = before.lower(), after.lower()
        blocked_by = None
        for ph in phrases:
            pl = ph.lower()
            if pl in bl and pl not in al:
                blocked_by = ph
                break
        if blocked_by:
            held.append({**e, "_hold_phrase": blocked_by})
        else:
            safe.append(e)
    return safe, held


def _normalize_issues(raw: Any) -> list[dict[str, Any]]:
    if not isinstance(raw, list):
        return []
    out: list[dict[str, Any]] = []
    for item in raw:
        if isinstance(item, dict):
            out.append(item)
    return out


def _issue_comment_body(issue: dict[str, Any]) -> str:
    cat = str(issue.get("category", "")).strip()
    sev = str(issue.get("severity", "")).strip()
    title = str(issue.get("title", "")).strip()
    expl = str(issue.get("explanation", "") or issue.get("details", "")).strip()
    why = str(issue.get("evidence_basis", "") or "").strip() or "(not specified — add evidence in a future pass if needed)"
    fix = str(issue.get("suggested_fix", "")).strip()
    conf = issue.get("confidence", "")
    agents = issue.get("source_agents") or []
    if not isinstance(agents, list):
        agents = [str(agents)]
    agents_s = ", ".join(str(a) for a in agents if str(a).strip()) or "(none)"
    posture = str(issue.get("accuracy_posture", "") or "").strip()
    lines = [
        f"[{cat}] [{sev}]",
        "",
        f"Title: {title}",
        "",
        "Problem:",
        expl or "(none)",
        "",
        "Why it matters:",
        why,
        "",
        "Suggested fix:",
        fix or "(none)",
        "",
        f"Confidence: {conf}",
        "",
        f"Source agents: {agents_s}",
    ]
    if posture:
        lines.extend(["", f"Accuracy posture: {posture}"])
    return "\n".join(lines).strip()


def _flatten_doc_like_extract(document_text: str) -> str:
    """Normalize newlines for loose comparison with python-docx flatten."""
    return re.sub(r"\r\n?", "\n", document_text or "").strip()


_PDF_PAGE_MARKER_LINE = re.compile(r"^---\s*PDF\s+page\s+\d+\s*---\s*$", re.IGNORECASE)


def strip_pdf_page_markers_for_display(text: str) -> str:
    """Remove PDF extraction page markers from user-facing exports (keeps body readable)."""
    if not text:
        return ""
    lines_out: list[str] = []
    for line in text.splitlines():
        if _PDF_PAGE_MARKER_LINE.match(line.strip()):
            continue
        lines_out.append(line.rstrip())
    cleaned = "\n".join(lines_out)
    cleaned = re.sub(r"\n{4,}", "\n\n\n", cleaned)
    return cleaned.strip()


def body_paragraphs_from_cleaned_plaintext(cleaned: str, *, max_paragraph_chars: int = 1200) -> list[str]:
    """Split normalized plaintext into Word-sized paragraphs (no fixed-width dumps)."""
    if not cleaned:
        return []
    parts = [p.strip() for p in cleaned.split("\n\n") if p.strip()]
    if len(parts) <= 1 and "\n" in cleaned:
        parts = [p.strip() for p in cleaned.split("\n") if p.strip()]
    out: list[str] = []
    for ch in parts:
        if len(ch) <= max_paragraph_chars:
            out.append(ch)
            continue
        start = 0
        while start < len(ch):
            slice_ = ch[start : start + max_paragraph_chars]
            br = slice_.rfind(". ")
            if br > max_paragraph_chars // 3:
                seg = ch[start : start + br + 1].strip()
                start = start + br + 2
            else:
                seg = slice_.strip()
                start += max_paragraph_chars
            if seg:
                out.append(seg)
    return out or [cleaned[:max_paragraph_chars]]


def _add_standard_review_footer(doc: Document, disc: DisclaimerBundle) -> None:
    doc.add_paragraph("")
    p = doc.add_paragraph()
    r = p.add_run("DraftLens — highlights and comments are editorial aids, not native Word Track Changes. ")
    r.bold = True
    r.font.size = Pt(10)
    r2 = p.add_run(disc.general)
    r2.font.size = Pt(10)
    p2 = doc.add_paragraph()
    r3 = p2.add_run("Structured issue list: see issues.pdf (and advanced artifacts for machine-readable exports).")
    r3.italic = True
    r3.font.size = Pt(9)


def _build_plaintext_review_docx(
    *,
    document_text: str,
    issues: list[dict[str, Any]],
    title: str,
    stamp: str,
    executive_summary: str,
    disc: DisclaimerBundle,
    format_notice: str | None,
) -> Document:
    """Build a readable reviewed .docx from normalized plaintext (PDF path or DOCX fallback)."""
    cleaned = strip_pdf_page_markers_for_display(document_text)
    doc = Document()
    doc.add_heading(title, level=1)
    doc.add_paragraph(f"Generated {stamp}")
    if format_notice:
        pn = doc.add_paragraph()
        rn = pn.add_run(format_notice)
        rn.italic = True
        rn.font.size = Pt(10)
    doc.add_paragraph(executive_summary or "Review completed.")
    doc.add_paragraph(disc.general)
    doc.add_heading("Manuscript", level=2)
    for para_text in body_paragraphs_from_cleaned_plaintext(cleaned):
        doc.add_paragraph(para_text)
    _decorate_review_docx(doc, issues, document_text=cleaned)
    _add_standard_review_footer(doc, disc)
    return doc


def _locate_issue_segments(
    issue: dict[str, Any],
    *,
    document_text: str,
    flat: str,
    spans: list[tuple[Any, int, int]],
) -> list[tuple[Any, int, int]]:
    """Return (paragraph, local_start, local_end) segments for highlighting/commenting."""
    span_text = str(issue.get("span_text") or "")
    cs = int(issue.get("char_start", -1))
    ce = int(issue.get("char_end", -1))
    dt = _flatten_doc_like_extract(document_text)
    fl = _flatten_doc_like_extract(flat)
    if span_text and 0 <= cs < ce <= len(dt) and dt == fl:
        frag = dt[cs:ce]
        if frag.strip() == span_text.strip():
            return dx.map_flat_range_to_paragraphs(spans, cs, ce)
    # Fallback: first paragraph containing span_text
    for p, _ps, _pe in spans:
        loc = dx.find_span_text_in_paragraph(p, span_text)
        if loc:
            return [(p, loc[0], loc[1])]
    return []


def _decorate_review_docx(
    doc: Document,
    issues: list[dict[str, Any]],
    *,
    document_text: str,
) -> None:
    for issue in issues:
        flat, spans = dx.flatten_doc_paragraphs(doc)
        segs = _locate_issue_segments(issue, document_text=document_text, flat=flat, spans=spans)
        if not segs:
            continue
        color = dx.severity_highlight(str(issue.get("severity", "minor")))
        body = _issue_comment_body(issue)
        for idx, (para, ls, le) in enumerate(segs):
            isolated = dx.isolate_range_as_runs(para, ls, le)
            if not isolated:
                continue
            first_r, last_r = isolated
            dx.highlight_run_span(first_r, last_r, color)
            if idx == 0:
                dx.try_anchor_comment(doc, para, first_r, last_r, body)


def _build_redline_html(
    *,
    document_text: str,
    issues: list[dict[str, Any]],
    disclaimer: DisclaimerBundle | None,
) -> str:
    excerpt = strip_pdf_page_markers_for_display((document_text or "")[:14_000])
    parts: list[str] = []
    for it in issues:
        cat = dx.html_escape(str(it.get("category", "")))
        sev = dx.html_escape(str(it.get("severity", "")))
        title = dx.html_escape(str(it.get("title", "")))
        span = dx.html_escape(str(it.get("span_text", "")))
        fix = dx.html_escape(str(it.get("suggested_fix", "")))
        expl = dx.html_escape(dx.snippet_for_display(str(it.get("explanation", "")), 400))
        parts.append(
            f'<article class="issue"><header><span class="pill cat">{cat}</span> '
            f'<span class="pill sev">{sev}</span> <strong>{title}</strong></header>'
            f'<p class="expl">{expl}</p>'
            f'<div class="diff"><span class="lbl">Before</span><del>{span}</del>'
            f'<span class="lbl">After</span><ins>{fix}</ins></div></article>'
        )
    body_inner = "\n".join(parts) if parts else '<p class="empty">No structured issues to diff.</p>'
    note = ""
    if disclaimer and disclaimer.third_party_models:
        note = f'<p class="foot">{dx.html_escape(disclaimer.third_party_models)}</p>'
    return f"""<!doctype html>
<html lang="en">
<head>
<meta charset="utf-8"/>
<meta name="viewport" content="width=device-width, initial-scale=1"/>
<title>DraftLens redline</title>
<style>
body {{ font-family: ui-sans-serif, system-ui, -apple-system, Segoe UI, Roboto, Helvetica, Arial; margin: 32px;
  color: #111; line-height: 1.45; max-width: 920px; }}
header {{ margin-bottom: 20px; }}
.muted {{ color: #555; font-size: 13px; }}
.excerpt {{ white-space: pre-wrap; background: #fafafa; border: 1px solid #e5e5e5; border-radius: 8px;
  padding: 14px; font-size: 13px; color: #333; margin-bottom: 28px; }}
.issue {{ border: 1px solid #e8e8e8; border-radius: 10px; padding: 14px 16px; margin-bottom: 16px;
  background: #fff; }}
.pill {{ display: inline-block; padding: 2px 8px; border-radius: 999px; font-size: 11px; font-weight: 600;
  margin-right: 6px; }}
.pill.cat {{ background: #eef2ff; color: #334155; }}
.pill.sev {{ background: #fff7ed; color: #9a3412; }}
.expl {{ font-size: 13px; color: #444; }}
.diff {{ margin-top: 10px; display: grid; gap: 8px; }}
.lbl {{ font-size: 11px; text-transform: uppercase; letter-spacing: 0.06em; color: #666; }}
del {{ background: #ffe4e6; color: #881337; padding: 2px 4px; border-radius: 4px; text-decoration: line-through; }}
ins {{ background: #dcfce7; color: #14532d; padding: 2px 4px; border-radius: 4px; text-decoration: none; }}
.foot {{ font-size: 12px; color: #666; margin-top: 24px; }}
.notice {{ font-size: 12px; color: #666; margin-top: 8px; }}
</style>
</head>
<body>
<header>
  <h1 style="font-size:22px;margin:0;">DraftLens redline</h1>
  <p class="muted">Generated {_now_stamp()} · Not native Word Track Changes</p>
  <p class="notice">Inline view below is generated from structured findings. Verify against the reviewed .docx.</p>
</header>
<section class="excerpt" aria-label="Document excerpt">{dx.html_escape(excerpt)}</section>
<section aria-label="Issues">{body_inner}</section>
{note}
</body>
</html>
"""


def _write_issues_markdown(issues: list[dict[str, Any]], path: Path) -> None:
    lines = ["# Issues", ""]
    for i, item in enumerate(issues, start=1):
        sev = str(item.get("severity", "note"))
        cat = str(item.get("category", ""))
        summ = str(item.get("title") or item.get("summary", "")).strip()
        det = str(item.get("explanation") or item.get("details", "")).strip()
        lines.append(f"{i}. **[{cat}] [{sev}]** — {summ}")
        if det:
            lines.extend(["", det, ""])
    path.write_text("\n".join(lines).strip() + "\n", encoding="utf-8")


def write_review_bundle(
    *,
    artifacts_dir: Path,
    original_docx_path: Path | None,
    document_text: str,
    arbiter_payload: dict[str, Any],
    debate_digest: str,
    has_supporting_files: bool = False,
    disclaimer_bundle: dict[str, Any] | None = None,
    sensitive_mode: bool = False,
    main_source_format: str = "docx",
    main_original_filename: str = "document",
) -> list[dict[str, str]]:
    artifacts_dir.mkdir(parents=True, exist_ok=True)
    disc = (
        DisclaimerBundle.model_validate(disclaimer_bundle)
        if disclaimer_bundle
        else build_disclaimer_bundle(
            sensitive_mode=sensitive_mode, has_supporting_files=has_supporting_files
        )
    )

    issues = _normalize_issues(arbiter_payload.get("issues"))
    if not issues:
        issues = [
            {
                "issue_id": "placeholder",
                "block_id": "doc-root",
                "severity": "minor",
                "category": "clarity",
                "title": "Structured issues missing",
                "explanation": "The model output did not include a structured issues array; see review_summary.md.",
                "suggested_fix": "",
                "confidence": 0.0,
                "source_agents": [],
                "span_text": "",
                "char_start": 0,
                "char_end": 0,
            }
        ]

    issues_path = artifacts_dir / "issues.json"
    issues_path.write_text(json.dumps(issues, indent=2, ensure_ascii=False) + "\n", encoding="utf-8")
    _write_issues_markdown(issues, artifacts_dir / "issues.md")

    summary = str(arbiter_payload.get("executive_summary", "")).strip()
    if not summary:
        summary = "Review completed. See issues.pdf and review_summary.pdf."
    footer = disc.markdown_review_footer.strip() or disc.general
    review_md = (
        f"# Review summary\n\n{_now_stamp()}\n\n{summary}\n\n## Debate digest\n\n{debate_digest.strip()}\n\n---\n\n{footer}\n"
    )
    (artifacts_dir / "review_summary.md").write_text(review_md, encoding="utf-8")

    redline_html = _build_redline_html(document_text=document_text, issues=issues, disclaimer=disc)
    (artifacts_dir / "redline.html").write_text(redline_html, encoding="utf-8")

    raw_dt = document_text or ""
    has_pdf_markers = bool(re.search(r"(?m)^---\s*PDF\s+page\s+\d+\s*---\s*$", raw_dt))
    decorate_text = strip_pdf_page_markers_for_display(raw_dt) if has_pdf_markers else raw_dt
    display_name = (main_original_filename or "document").strip() or "document"
    title = f"Reviewed: {display_name}"
    fmt_notice: str | None = None
    if str(main_source_format).lower() == "pdf":
        fmt_notice = (
            "Source was PDF — layout is approximate and based on extracted text. "
            "Highlights and comments mark findings in the readable export below."
        )

    reviewed_path = artifacts_dir / "reviewed.docx"
    if original_docx_path and original_docx_path.exists():
        shutil.copy2(original_docx_path, reviewed_path)
        try:
            doc = Document(str(reviewed_path))
            _decorate_review_docx(doc, issues, document_text=decorate_text)
            _add_standard_review_footer(doc, disc)
            doc.save(str(reviewed_path))
        except Exception:
            logger.exception("reviewed docx pipeline failed; emitting structured plaintext export")
            doc = _build_plaintext_review_docx(
                document_text=document_text,
                issues=issues,
                title=f"Reviewed export (fallback) — {display_name}",
                stamp=_now_stamp(),
                executive_summary=summary
                + "\n\nThe source DOCX could not be decorated in place (complex or non-standard OOXML). "
                "This document rebuilds readable body text without PDF page markers.",
                disc=disc,
                format_notice=fmt_notice,
            )
            doc.save(str(reviewed_path))
    else:
        doc = _build_plaintext_review_docx(
            document_text=document_text,
            issues=issues,
            title=title,
            stamp=_now_stamp(),
            executive_summary=summary,
            disc=disc,
            format_notice=fmt_notice,
        )
        doc.save(str(reviewed_path))

    return [
        {"name": "reviewed.docx", "path": str(reviewed_path), "media_type": reviewed_docx_media()},
        {"name": "review_summary.md", "path": str(artifacts_dir / "review_summary.md"), "media_type": "text/markdown"},
        {"name": "issues.md", "path": str(artifacts_dir / "issues.md"), "media_type": "text/markdown"},
        {"name": "issues.json", "path": str(issues_path), "media_type": "application/json"},
        {"name": "redline.html", "path": str(artifacts_dir / "redline.html"), "media_type": "text/html"},
    ]


def reviewed_docx_media() -> str:
    return "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


def _normalize_change_rows(
    issues: list[dict[str, Any]], changes_raw: Any
) -> list[dict[str, Any]]:
    rows: list[dict[str, Any]] = []
    if not isinstance(changes_raw, list):
        return rows
    issue_by_block: dict[str, list[dict[str, Any]]] = {}
    for it in issues:
        bid = str(it.get("block_id", ""))
        issue_by_block.setdefault(bid, []).append(it)

    for ch in changes_raw:
        if not isinstance(ch, dict):
            continue
        before = str(ch.get("before", ""))
        after = str(ch.get("after", ""))
        rationale = str(ch.get("rationale", "")).strip()
        edit_id = str(ch.get("edit_id", "") or "")
        block_id = str(ch.get("block_id", ""))
        matched = None
        for it in issue_by_block.get(block_id, []):
            if before and str(it.get("span_text", "")) == before:
                matched = it
                break
        if matched is None and before:
            for it in issues:
                if str(it.get("span_text", "")) == before:
                    matched = it
                    break
        issue_id = str((matched or {}).get("issue_id", edit_id or "unknown"))
        cat = str((matched or {}).get("category", ch.get("category", "")))
        sev = str((matched or {}).get("severity", ch.get("severity", "")))
        rows.append(
            {
                "issue_id": issue_id,
                "category": cat,
                "severity": sev,
                "original_snippet": before,
                "revised_snippet": after,
                "rationale": rationale,
                "block_id": block_id,
                "edit_id": edit_id,
            }
        )
    return rows


def _apply_edits_to_docx(document: Document, edits: list[dict[str, Any]]) -> int:
    """Apply in-paragraph replacements; returns count of paragraphs modified."""
    ordered = sorted(
        [e for e in edits if isinstance(e, dict) and str(e.get("before", "")).strip()],
        key=lambda e: len(str(e.get("before", ""))),
        reverse=True,
    )
    changed = 0
    for para in dx.iter_body_paragraphs(document):
        txt = para.text or ""
        new_txt = txt
        for e in ordered:
            b = str(e.get("before", ""))
            a = str(e.get("after", ""))
            if b and b in new_txt:
                new_txt = new_txt.replace(b, a, 1)
        if new_txt != txt:
            dx.set_paragraph_text_unicode_safe(para, new_txt)
            changed += 1
    return changed


def _apply_full_corrected_text(document: Document, corrected_text: str) -> None:
    """Replace body paragraph texts line-wise (layout best-effort fallback)."""
    lines = corrected_text.splitlines()
    paras = list(dx.iter_body_paragraphs(document))
    for i, p in enumerate(paras):
        dx.set_paragraph_text_unicode_safe(p, lines[i] if i < len(lines) else "")
    if len(lines) > len(paras):
        for j in range(len(paras), len(lines)):
            document.add_paragraph(lines[j])


def write_fix_bundle(
    *,
    artifacts_dir: Path,
    original_docx_path: Path | None,
    document_text: str,
    arbiter_payload: dict[str, Any],
    has_supporting_files: bool = False,
    disclaimer_bundle: dict[str, Any] | None = None,
    sensitive_mode: bool = False,
    do_not_change: str | None = None,
) -> list[dict[str, str]]:
    artifacts_dir.mkdir(parents=True, exist_ok=True)
    disc = (
        DisclaimerBundle.model_validate(disclaimer_bundle)
        if disclaimer_bundle
        else build_disclaimer_bundle(
            sensitive_mode=sensitive_mode, has_supporting_files=has_supporting_files
        )
    )

    issues = _normalize_issues(arbiter_payload.get("issues"))
    changes_raw = arbiter_payload.get("changes")
    changes_for_doc = [c for c in (changes_raw if isinstance(changes_raw, list) else []) if isinstance(c, dict)]
    phrases = _locked_phrases(do_not_change)
    safe_edits, held_edits = _partition_edits_for_do_not_change(changes_for_doc, phrases)

    change_rows = _normalize_change_rows(issues, safe_edits)
    for h in held_edits:
        phrase = str(h.get("_hold_phrase", ""))
        change_rows.append(
            {
                "issue_id": "do-not-change-hold",
                "kind": "do_not_change_hold",
                "category": "risk",
                "severity": "major",
                "original_snippet": str(h.get("before", "")),
                "revised_snippet": str(h.get("after", "")),
                "rationale": (
                    "Not applied automatically: the edit would alter or remove a phrase listed under "
                    f"do-not-change (`{phrase}`). Preserve meaning; apply manually only if policy allows."
                ),
                "block_id": str(h.get("block_id", "")),
                "edit_id": str(h.get("edit_id", "")),
            }
        )

    if not change_rows:
        change_rows = [
            {
                "issue_id": "n/a",
                "category": "clarity",
                "severity": "minor",
                "original_snippet": "",
                "revised_snippet": "",
                "rationale": "No structured changes returned; see corrected body if present.",
                "block_id": "",
                "edit_id": "",
            }
        ]

    (artifacts_dir / "changes.json").write_text(
        json.dumps(change_rows, indent=2, ensure_ascii=False) + "\n", encoding="utf-8"
    )

    lines = ["# Change log", "", f"Generated {_now_stamp()}", ""]
    for i, ch in enumerate(change_rows, start=1):
        lines.append(f"## {i}. {ch.get('issue_id')}")
        if ch.get("kind") == "do_not_change_hold":
            lines.append("")
            lines.append("_Held: do-not-change safeguard._")
        lines.append("")
        lines.append(f"- **category:** {ch.get('category')}")
        lines.append(f"- **severity:** {ch.get('severity')}")
        lines.append("")
        lines.append("**original_snippet**")
        lines.append("")
        lines.append(str(ch.get("original_snippet", "")))
        lines.append("")
        lines.append("**revised_snippet**")
        lines.append("")
        lines.append(str(ch.get("revised_snippet", "")))
        lines.append("")
        lines.append("**rationale**")
        lines.append("")
        lines.append(str(ch.get("rationale", "")))
        lines.append("")
    lines.append("---")
    lines.append("")
    foot = disc.markdown_fix_footer.strip() or disc.retention or ""
    lines.append(foot)
    if disc.no_model_training:
        lines.append("")
        lines.append(str(disc.no_model_training))
    lines.append("")
    (artifacts_dir / "change_log.md").write_text("\n".join(lines).strip() + "\n", encoding="utf-8")

    corrected_path = artifacts_dir / "corrected.docx"
    corrected_text = str(arbiter_payload.get("corrected_document_text", "")).strip()

    if original_docx_path and original_docx_path.exists():
        shutil.copy2(original_docx_path, corrected_path)
        try:
            doc = Document(str(corrected_path))
            n = _apply_edits_to_docx(doc, safe_edits)
            if n == 0 and corrected_text and not held_edits:
                _apply_full_corrected_text(doc, corrected_text)
            doc.save(str(corrected_path))
        except Exception:
            logger.exception("corrected docx pipeline failed; emitting text fallback")
            doc = Document()
            doc.add_heading("DraftLens corrected export (fallback)", level=1)
            doc.add_paragraph(f"Generated {_now_stamp()}")
            doc.add_paragraph(
                "The source DOCX could not be rewritten safely. See change_log.md for structured edits."
            )
            body = corrected_text or document_text
            cleaned_body = strip_pdf_page_markers_for_display(body)
            for para_text in body_paragraphs_from_cleaned_plaintext(cleaned_body):
                doc.add_paragraph(para_text)
            doc.save(str(corrected_path))
    else:
        doc = Document()
        doc.add_heading("DraftLens corrected export", level=1)
        doc.add_paragraph(f"Generated {_now_stamp()}")
        body = corrected_text or document_text
        cleaned_body = strip_pdf_page_markers_for_display(body)
        for para_text in body_paragraphs_from_cleaned_plaintext(cleaned_body):
            doc.add_paragraph(para_text)
        doc.save(str(corrected_path))

    return [
        {"name": "corrected.docx", "path": str(corrected_path), "media_type": reviewed_docx_media()},
        {"name": "change_log.md", "path": str(artifacts_dir / "change_log.md"), "media_type": "text/markdown"},
        {"name": "changes.json", "path": str(artifacts_dir / "changes.json"), "media_type": "application/json"},
    ]
