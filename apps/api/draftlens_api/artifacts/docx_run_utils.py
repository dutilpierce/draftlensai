from __future__ import annotations

import re
from typing import Any, Iterator

from docx.document import Document as DocumentObject
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph
from docx.text.run import Run


def set_paragraph_text_unicode_safe(paragraph: Paragraph, new_text: str) -> None:
    """
    Replace full paragraph text without assigning `paragraph.text`, which clears runs and
    drops per-run font hints — a common cause of CJK / complex-script corruption in Word.
    """
    text = new_text or ""
    runs = list(paragraph.runs)
    if not runs:
        paragraph.add_run(text)
        return
    template = runs[0]
    for r in runs[1:]:
        el = r._element
        parent = el.getparent()
        if parent is not None:
            parent.remove(el)
    template.text = text


def set_cell_text_unicode_safe(cell: _Cell, new_text: str) -> None:
    """Set first paragraph of a table cell without corrupting fonts."""
    if cell.paragraphs:
        set_paragraph_text_unicode_safe(cell.paragraphs[0], new_text)
    else:
        cell.add_paragraph(new_text or "")


def iter_body_paragraphs(parent: Any) -> Iterator[Paragraph]:
    """
    Walk main document body and table cells (not headers/footers).

    Headers/footers are skipped because python-docx does not expose them on Document
    the same way; comment anchoring there is unsupported for this renderer.
    """
    if isinstance(parent, DocumentObject):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        return
    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            table = Table(child, parent)
            for row in table.rows:
                for cell in row.cells:
                    yield from iter_body_paragraphs(cell)


def flatten_doc_paragraphs(document: DocumentObject) -> tuple[str, list[tuple[Paragraph, int, int]]]:
    """
    Build a newline-joined plain text view of body paragraphs and (para, start, end) offsets.

    Offsets are half-open [start, end) into the flat string (including \\n separators).
    """
    parts: list[str] = []
    spans: list[tuple[Paragraph, int, int]] = []
    pos = 0
    first = True
    for p in iter_body_paragraphs(document):
        txt = p.text or ""
        if not first:
            parts.append("\n")
            pos += 1
        spans.append((p, pos, pos + len(txt)))
        parts.append(txt)
        pos += len(txt)
        first = False
    return "".join(parts), spans


def _split_run(run: Run, split_index: int) -> tuple[Run, Run]:
    """Split `run` into two runs at character offset `split_index` (0 < i < len)."""
    text = run.text or ""
    if split_index <= 0 or split_index >= len(text):
        return run, run
    head, tail = text[:split_index], text[split_index:]
    run.text = head
    tail_run = run._parent.add_run(tail)  # type: ignore[attr-defined]
    run._r.addnext(tail_run._r)
    return run, tail_run


def isolate_range_as_runs(paragraph: Paragraph, start: int, end: int) -> tuple[Run, Run] | None:
    """
    Split paragraph runs so that text[start:end] is covered by contiguous runs.

    Returns (first_run, last_run) inclusive for the isolated range, or None if invalid.
    """
    full = paragraph.text or ""
    if start < 0 or end > len(full) or start >= end:
        return None

    runs = list(paragraph.runs)
    if not runs:
        return None

    # Map run boundaries in paragraph text
    boundaries: list[tuple[Run, int, int]] = []
    cursor = 0
    for r in runs:
        t = r.text or ""
        ln = len(t)
        boundaries.append((r, cursor, cursor + ln))
        cursor += ln

    if cursor != len(full):
        # Rare mismatch (field codes etc.); fall back to single-run paragraph rebuild
        return None

    # Split at start and end boundaries
    for r, rs, rend in boundaries:
        if rs <= start < rend:
            if start > rs:
                _split_run(r, start - rs)
            break
    # Recompute after first split
    runs = list(paragraph.runs)
    boundaries = []
    cursor = 0
    for r in runs:
        t = r.text or ""
        ln = len(t)
        boundaries.append((r, cursor, cursor + ln))
        cursor += ln

    for r, rs, rend in boundaries:
        if rs < end <= rend:
            if end < rend:
                _split_run(r, end - rs)
            break

    runs = list(paragraph.runs)
    boundaries = []
    cursor = 0
    for r in runs:
        t = r.text or ""
        ln = len(t)
        boundaries.append((r, cursor, cursor + ln))
        cursor += ln

    first_run: Run | None = None
    last_run: Run | None = None
    for r, rs, rend in boundaries:
        if rs < end and rend > start:
            if first_run is None:
                first_run = r
            last_run = r
    if first_run is None or last_run is None:
        return None
    return first_run, last_run


def map_flat_range_to_paragraphs(
    spans: list[tuple[Paragraph, int, int]], start: int, end: int
) -> list[tuple[Paragraph, int, int]]:
    """Return list of (paragraph, local_start, local_end) slices intersecting [start, end)."""
    out: list[tuple[Paragraph, int, int]] = []
    for p, ps, pe in spans:
        if pe <= start or ps >= end:
            continue
        ls = max(0, start - ps)
        le = min(pe - ps, end - ps)
        out.append((p, ls, le))
    return out


def severity_highlight(severity: str) -> WD_COLOR_INDEX:
    s = (severity or "minor").lower()
    if s == "critical":
        return WD_COLOR_INDEX.RED
    if s == "major":
        return WD_COLOR_INDEX.YELLOW
    if s == "minor":
        return WD_COLOR_INDEX.TURQUOISE
    return WD_COLOR_INDEX.GRAY_25


def highlight_run_span(first_run: Run, last_run: Run, color: WD_COLOR_INDEX) -> None:
    """Apply highlight across contiguous runs in same paragraph."""
    collecting = False
    for r in first_run._parent.runs:  # type: ignore[attr-defined]
        if r is first_run:
            collecting = True
        if collecting:
            r.font.highlight_color = color
        if r is last_run:
            break


def try_anchor_comment(
    document: DocumentObject,
    paragraph: Paragraph,
    first_run: Run,
    last_run: Run,
    comment_text: str,
) -> bool:
    """
    Add a native Word comment anchored to [first_run, last_run].

    Returns False if comments cannot be anchored (should be rare).
    """
    try:
        comments = document.part.comments
        c = comments.add_comment(text=comment_text, author="DraftLens", initials="DL")
        first_run.mark_comment_range(last_run, c.comment_id)
        return True
    except Exception:  # noqa: BLE001
        return False


def find_span_text_in_paragraph(paragraph: Paragraph, needle: str) -> tuple[int, int] | None:
    """Return (start, end) indices within paragraph.text for first occurrence of needle."""
    hay = paragraph.text or ""
    if not needle.strip():
        return None
    i = hay.find(needle)
    if i < 0:
        return None
    return i, i + len(needle)


def escape_xml_text(s: str) -> str:
    return (
        (s or "")
        .replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
        .replace('"', "&quot;")
    )


def html_escape(s: str) -> str:
    return escape_xml_text(s)


def snippet_for_display(s: str, max_len: int = 280) -> str:
    t = re.sub(r"\s+", " ", (s or "").strip())
    if len(t) <= max_len:
        return t
    return t[: max_len - 1] + "…"
