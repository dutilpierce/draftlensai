"""
Unified text extraction for main manuscripts and supporting evidence.

All supported formats normalize to plain text plus page metadata where applicable.
OCR is not implemented: image-only PDFs fail the main path with a clear message.
"""

from __future__ import annotations

import re
import zipfile
from dataclasses import dataclass
from pathlib import Path
from typing import Literal

from docx import Document
from pypdf import PdfReader

_DOC_NOT_SUPPORTED_MSG = (
    "DOC files are not fully supported yet. Please upload DOCX or PDF."
)

_PDF_NO_TEXT_MSG = (
    "This PDF has no selectable text (often scanned or image-only pages). "
    "DraftLens does not run OCR yet. Please upload a searchable PDF, export text, or use DOCX."
)

_PDF_LOW_TEXT_MSG = (
    "This PDF has very little selectable text relative to its page count (often scans, forms, or heavy layout). "
    "DraftLens does not run OCR yet. Try a text-based export or DOCX."
)


class UserFacingDocumentError(ValueError):
    """Raised for main-manuscript problems that should surface as HTTP 400 with `str(exc)`."""


@dataclass(frozen=True)
class ParsedPage:
    """One-based page index and extracted text for that page."""

    page_number: int
    text: str


@dataclass(frozen=True)
class NormalizedDocument:
    """Shared extraction shape for main and supporting ingestion."""

    plain_text: str
    page_count: int
    mime: str
    source_format: Literal["docx", "pdf", "doc", "txt", "md", "unknown"]
    parse_notes: tuple[str, ...] = ()
    extraction_quality: Literal["ok", "low_text", "empty"] = "ok"
    pages: tuple[ParsedPage, ...] = ()


@dataclass(frozen=True)
class MainExtractionResult:
    """Result of parsing the primary review target (DOCX or PDF)."""

    text: str
    pages: int
    mime: str
    source_format: Literal["docx", "pdf"]
    parse_notes: tuple[str, ...] = ()


def estimate_docx_pages(doc: Document) -> int:
    text = "\n".join(p.text for p in doc.paragraphs if p.text)
    words = len(re.findall(r"\w+", text))
    pages = max(1, int(words / 280) + (1 if words % 280 else 0))
    return min(pages, 5000)


def extract_docx(path: Path) -> tuple[str, int]:
    if not path.is_file():
        raise ValueError("docx_missing")
    if not zipfile.is_zipfile(path):
        raise ValueError("docx_invalid_or_corrupt")
    try:
        document = Document(str(path))
    except (KeyError, OSError, ValueError) as exc:
        raise ValueError(f"docx_open_failed:{exc}") from exc
    parts: list[str] = []
    for p in document.paragraphs:
        if p.text and p.text.strip():
            parts.append(p.text.strip())
    return "\n\n".join(parts), estimate_docx_pages(document)


def extract_plain(path: Path) -> tuple[str, int]:
    raw = path.read_text(encoding="utf-8", errors="replace")
    lines = [ln.strip() for ln in raw.splitlines() if ln.strip()]
    body = "\n".join(lines)
    words = len(re.findall(r"\w+", body))
    pages = max(1, int(words / 350) + (1 if words % 350 else 0))
    return body, pages


def _pdf_quality(chars: int, page_count: int, words: int) -> Literal["ok", "low_text", "empty"]:
    if words < 8 or chars < 12:
        return "empty"
    threshold_chars = max(140, page_count * 55)
    threshold_words = max(40, page_count * 22)
    if chars < threshold_chars and words < threshold_words:
        return "low_text"
    return "ok"


def extract_pdf_normalized(path: Path, *, strict: bool) -> NormalizedDocument:
    """Extract PDF text with page markers; `strict` raises on empty/low-text for main manuscripts."""
    if not path.is_file():
        raise ValueError("pdf_missing")
    reader = PdfReader(str(path))
    page_count = max(1, len(reader.pages))
    parsed_pages: list[ParsedPage] = []
    segments: list[str] = []
    for i, page in enumerate(reader.pages):
        raw = page.extract_text() or ""
        t = raw.strip()
        parsed_pages.append(ParsedPage(page_number=i + 1, text=t))
        if t:
            segments.append(f"--- PDF page {i + 1} ---\n\n{t}")
    body = "\n\n".join(segments)
    words = len(re.findall(r"\w+", body))
    chars = len(body.strip())
    quality = _pdf_quality(chars, page_count, words)
    notes: list[str] = []
    if quality == "low_text":
        notes.append("Low selectable text relative to page count.")
    if quality == "empty":
        notes.append("No selectable text extracted from PDF pages.")

    if strict and quality == "empty":
        raise UserFacingDocumentError(_PDF_NO_TEXT_MSG)
    if strict and quality == "low_text":
        raise UserFacingDocumentError(_PDF_LOW_TEXT_MSG)

    return NormalizedDocument(
        plain_text=body,
        page_count=page_count,
        mime="application/pdf",
        source_format="pdf",
        parse_notes=tuple(notes),
        extraction_quality=quality,
        pages=tuple(parsed_pages),
    )


def extract_pdf(path: Path) -> tuple[str, int]:
    """Backward-compatible PDF extraction (page markers, same quality rules as supporting path)."""
    norm = extract_pdf_normalized(path, strict=False)
    return norm.plain_text, norm.page_count


def parse_main_manuscript(path: Path, original_filename: str) -> MainExtractionResult:
    """Parse the uploaded main review document (DOCX or PDF). Raises UserFacingDocumentError for bad PDFs."""
    _ = original_filename
    suf = path.suffix.lower()
    if suf == ".docx":
        text, pages = extract_docx(path)
        notes: tuple[str, ...] = ()
        if not text.strip():
            notes = ("DOCX contained no paragraph text.",)
        return MainExtractionResult(
            text=text,
            pages=max(1, pages),
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            source_format="docx",
            parse_notes=notes,
        )
    if suf == ".pdf":
        norm = extract_pdf_normalized(path, strict=True)
        notes_list = list(norm.parse_notes)
        notes_list.append("Main document ingested from PDF; page markers are included in the extracted text.")
        return MainExtractionResult(
            text=norm.plain_text,
            pages=norm.page_count,
            mime="application/pdf",
            source_format="pdf",
            parse_notes=tuple(notes_list),
        )
    raise ValueError("unsupported_main_document")


def normalize_supporting_file(path: Path, original_name: str) -> NormalizedDocument:
    """
    Evidence-only normalization. Never raises for low-text PDF (partial quality);
    `.doc` yields an explicit placeholder (no silent success).
    """
    suffix = path.suffix.lower()
    if suffix == ".docx":
        text, pages = extract_docx(path)
        return NormalizedDocument(
            plain_text=text,
            page_count=pages,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            source_format="docx",
            parse_notes=(),
            extraction_quality="ok" if text.strip() else "empty",
            pages=(),
        )
    if suffix == ".pdf":
        return extract_pdf_normalized(path, strict=False)
    if suffix in {".txt", ".md"}:
        text, pages = extract_plain(path)
        return NormalizedDocument(
            plain_text=text,
            page_count=pages,
            mime="text/plain",
            source_format="txt" if suffix == ".txt" else "md",
            parse_notes=(),
            extraction_quality="ok" if text.strip() else "empty",
            pages=(),
        )
    if suffix == ".doc":
        return NormalizedDocument(
            plain_text=_DOC_NOT_SUPPORTED_MSG,
            page_count=1,
            mime="application/msword",
            source_format="doc",
            parse_notes=("doc_not_supported",),
            extraction_quality="empty",
            pages=(),
        )
    raise ValueError(f"unsupported_supporting:{suffix}")


def extract_supporting(path: Path, filename: str) -> str:
    """Plain string extraction for simple bundle paths (e.g. evidence/ingestion.py)."""
    return normalize_supporting_file(path, filename).plain_text
